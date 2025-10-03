import re
from typing import Dict, Set
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

PLACEHOLDER_RE = re.compile(r"\{[A-Z0-9_]+\}")


def find_placeholders_in_text(text: str) -> Set[str]:
    return set(PLACEHOLDER_RE.findall(text or ""))


def iter_paragraphs_and_cells(doc: Document):
    # Параграфы
    for p in doc.paragraphs:
        yield p
    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_placeholders(doc: Document, mapping: Dict[str, str]):
    """ Аккуратно заменяет плейсхолдеры, с учетом разбиения на runs. """
    for p in iter_paragraphs_and_cells(doc):
        replace_in_paragraph(p, mapping)


def replace_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]):
    # Склеиваем все runs в один текст, запоминаем границы
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return
    placeholders = find_placeholders_in_text(full_text)
    if not placeholders:
        return
    # Замена всех найденных
    for ph in placeholders:
        if ph in mapping:
            full_text = full_text.replace(ph, mapping[ph])
    # Переписываем в один run
    for _ in range(len(paragraph.runs) - 1):
        paragraph.runs[0].merge(paragraph.runs[1])
    paragraph.runs[0].text = full_text


def extract_placeholders(doc: Document) -> Set[str]:
    found = set()
    for p in iter_paragraphs_and_cells(doc):
        found |= find_placeholders_in_text(p.text)
    return found